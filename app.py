# At the top of app.py
DEBUG = False  # Change to False in production

from flask import Flask, render_template, request, send_from_directory, flash, redirect, url_for
import os
import platform
import subprocess
import pandas as pd
from docx import Document
from werkzeug.utils import secure_filename
from datetime import datetime
from babel.dates import format_date

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'

# Configuration
app.config['EXCEL_FILE'] = 'data/data.xlsx'
app.config['TEMPLATE_DIR'] = 'data/templates'
app.config['OUTPUT_DIR'] = 'output_docs'
app.config['ALLOWED_EXTENSIONS'] = {'xlsx'}

# Helper functions
arabic_months = {
    "January": "كانون الثاني",
    "February": "شباط",
    "March": "آذار",
    "April": "نيسان",
    "May": "أيار",
    "June": "حزيران",
    "July": "تموز",
    "August": "آب",
    "September": "أيلول",
    "October": "تشرين الأول",
    "November": "تشرين الثاني",
    "December": "كانون الأول"
}

def convert_to_eastern_arabic(number):
    eastern_arabic_digits = {
        "0": "٠", "1": "١", "2": "٢", "3": "٣", "4": "٤",
        "5": "٥", "6": "٦", "7": "٧", "8": "٨", "9": "٩"
    }
    return "".join(eastern_arabic_digits[digit] for digit in str(number))

def clean_date_value(value):
    """Clean date values to remove unwanted time components"""
    if pd.isna(value) or value == '':
        return ''
    
    # If it's already a string and doesn't contain time info, return as is
    if isinstance(value, str):
        # Check if it looks like a date with time (contains time component)
        if ' 00:00:00' in value:
            return value.replace(' 00:00:00', '')
        return value
    
    # If it's a datetime object, format it as date only
    if isinstance(value, (pd.Timestamp, datetime)):
        return value.strftime('%Y-%m-%d')
    
    # If it's a date object
    if hasattr(value, 'strftime'):
        return value.strftime('%Y-%m-%d')
    
    # For any other type, convert to string and clean
    str_value = str(value)
    if ' 00:00:00' in str_value:
        return str_value.replace(' 00:00:00', '')
    
    return str_value

def clean_dataframe_dates(df):
    """Clean all date columns in the dataframe"""
    # List of common date column names (adjust based on your Excel structure)
    date_columns = ['date', 'Date', 'birth_date', 'Birth Date', 'baptism_date', 'Baptism Date', 
                   'marriage_date', 'Marriage Date', 'created_date', 'updated_date']
    
    # Also check for columns that might contain dates based on data type
    for col in df.columns:
        # Check if column name suggests it's a date
        col_lower = col.lower()
        is_date_column = any(date_word in col_lower for date_word in ['date', 'birth', 'baptism', 'marriage', 'created', 'updated'])
        
        # Or check if the column contains datetime-like data
        if is_date_column or df[col].dtype == 'datetime64[ns]':
            df[col] = df[col].apply(clean_date_value)
    
    return df

def replace_text_in_paragraph(paragraph, old_text, new_text):
    """Replace text while preserving original formatting"""
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text in full_text:
        # Store original formatting from the first run that has text
        original_font = None
        for run in paragraph.runs:
            if run.text.strip():
                original_font = run.font
                break
        
        # Replace text
        full_text = full_text.replace(old_text, new_text)
        
        # Clear all runs
        for run in paragraph.runs:
            run.text = ""
        
        # Add new run with preserved formatting
        new_run = paragraph.add_run(full_text)
        if original_font:
            new_run.font.name = original_font.name
            new_run.font.size = original_font.size
            new_run.font.bold = original_font.bold
            new_run.font.italic = original_font.italic
            new_run.font.underline = original_font.underline
            if original_font.color and original_font.color.rgb:
                new_run.font.color.rgb = original_font.color.rgb

def replace_text_in_paragraph_advanced(paragraph, old_text, new_text):
    """Advanced text replacement that enforces a reliable Arabic/English font"""
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return
    
    # Find the run that contains the start of old_text
    char_index = 0
    target_run_index = -1
    
    for i, run in enumerate(paragraph.runs):
        if char_index <= full_text.find(old_text) < char_index + len(run.text):
            target_run_index = i
            break
        char_index += len(run.text)
    
    if target_run_index >= 0:
        target_run = paragraph.runs[target_run_index]
        
        # Store original formatting (but we'll force a reliable font)
        original_font_size = target_run.font.size
        original_bold = target_run.font.bold
        original_italic = target_run.font.italic
        original_underline = target_run.font.underline
        original_color = target_run.font.color.rgb if target_run.font.color and target_run.font.color.rgb else None
        
        # Replace using simple method
        full_text = full_text.replace(old_text, new_text)
        
        # Clear all runs
        for run in paragraph.runs:
            run.text = ""
        
        # Add new run with preserved formatting + reliable font
        new_run = paragraph.add_run(full_text)
        
        # Try fonts in order of preference (most reliable first)
        font_options = [
            'Tahoma',           # Excellent Arabic/English support, widely available
            'Arial Unicode MS', # Good fallback, supports Arabic well
            'Calibri',          # Modern, good Arabic support
            'Times New Roman',  # Classic, reliable Arabic support
            'Segoe UI',         # Modern Windows font with Arabic
            'Arial'             # Ultimate fallback
        ]
        
        # Use the first available font (Tahoma is usually the best choice)
        preferred_font = font_options[0]  # Tahoma
        new_run.font.name = preferred_font
        
        # Apply other preserved formatting
        if original_font_size:
            new_run.font.size = original_font_size
        if original_bold is not None:
            new_run.font.bold = original_bold
        if original_italic is not None:
            new_run.font.italic = original_italic
        if original_underline is not None:
            new_run.font.underline = original_underline
        if original_color:
            new_run.font.color.rgb = original_color
            
        # Set complex script font for Arabic text (important for RTL languages)
        new_run.font.cs_font = preferred_font
        
        print(f"Applied font: {preferred_font} to text: '{new_text[:30]}...'")

# Alternative function if you want to specify a different font
def replace_text_with_custom_font(paragraph, old_text, new_text, font_name='Times New Roman'):
    """Replace text with a specific font and ensure consistent formatting"""
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return
    
    # Store the original paragraph formatting
    paragraph_format = paragraph.paragraph_format
    
    # Get the first run's formatting as baseline
    baseline_run = None
    for run in paragraph.runs:
        if run.text.strip():
            baseline_run = run
            break
    
    if not baseline_run:
        baseline_run = paragraph.runs[0] if paragraph.runs else None
    
    # Store formatting properties
    if baseline_run:
        original_size = baseline_run.font.size
        original_bold = baseline_run.font.bold
        original_italic = baseline_run.font.italic
        original_underline = baseline_run.font.underline
        original_color = baseline_run.font.color.rgb if baseline_run.font.color and baseline_run.font.color.rgb else None
    else:
        # Default formatting
        from docx.shared import Pt
        original_size = Pt(12)
        original_bold = False
        original_italic = False
        original_underline = False
        original_color = None
    
    # Replace text
    full_text = full_text.replace(old_text, new_text)
    
    # Clear all runs
    for run in paragraph.runs[:]:
        paragraph._element.remove(run._element)
    
    # Create a single new run with consistent formatting
    new_run = paragraph.add_run(full_text)
    
    # Apply consistent font and formatting
    new_run.font.name = font_name
    new_run.font.cs_font = font_name  # Complex script font for Arabic
    new_run.font.east_asian_font = font_name  # East Asian font
    
    # Apply preserved formatting
    if original_size:
        new_run.font.size = original_size
    if original_bold is not None:
        new_run.font.bold = original_bold
    if original_italic is not None:
        new_run.font.italic = original_italic
    if original_underline is not None:
        new_run.font.underline = original_underline
    if original_color:
        new_run.font.color.rgb = original_color
    
    print(f"Applied consistent font '{font_name}' to: '{new_text[:50]}...'")

def apply_consistent_font_to_document(doc, font_name='Times New Roman'):
    """Apply consistent font to entire document after all replacements"""
    print(f"Applying consistent font '{font_name}' to entire document...")
    
    # Apply to all paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():  # Only apply to runs with actual text
                run.font.name = font_name
                run.font.cs_font = font_name
                run.font.east_asian_font = font_name
    
    # Apply to tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text.strip():
                            run.font.name = font_name
                            run.font.cs_font = font_name
                            run.font.east_asian_font = font_name
    
    # Apply to headers and footers
    for section in doc.sections:
        # Header
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                if run.text.strip():
                    run.font.name = font_name
                    run.font.cs_font = font_name
                    run.font.east_asian_font = font_name
        
        # Footer
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                if run.text.strip():
                    run.font.name = font_name
                    run.font.cs_font = font_name
                    run.font.east_asian_font = font_name
    
    print(f"Consistent font '{font_name}' applied to entire document.")

def fill_template(template_path, data_row, output_path):
    if not os.path.exists(template_path):
        return False, f"الملف غير موجود: {template_path}"
    
    try:
        doc = Document(template_path)
    except Exception as e:
        return False, f"لم استطع فتح الملف: {e}"
    
    # Set preferred font
    preferred_font = PREFERRED_ARABIC_FONT
    
    # Debug: Print available columns and placeholders
    print("=== TEMPLATE FILLING DEBUG ===")
    print("Available data columns:", list(data_row.keys()))
    print(f"Using font: {preferred_font}")
    
    # Find all placeholders in the document
    placeholders_found = set()
    for paragraph in doc.paragraphs:
        import re
        placeholders = re.findall(r'\{([^}]+)\}', paragraph.text)
        placeholders_found.update(placeholders)
    
    print("Placeholders found in template:", sorted(placeholders_found))
    
    # Replace placeholders while using reliable font
    replacements_made = 0
    for key, value in data_row.items():
        key = key.strip()
        clean_value = clean_date_value(value)
        placeholder = f'{{{key}}}'
        
        # Check if this placeholder exists in template
        placeholder_exists = False
        
        # Search in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                placeholder_exists = True
                print(f"Replacing '{placeholder}' with '{clean_value}' in paragraph")
                replace_text_with_custom_font(paragraph, placeholder, str(clean_value), preferred_font)
                replacements_made += 1
        
        # Search in tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            placeholder_exists = True
                            print(f"Replacing '{placeholder}' with '{clean_value}' in table")
                            replace_text_with_custom_font(paragraph, placeholder, str(clean_value), preferred_font)
                            replacements_made += 1
        
        # Search in headers and footers
        for section in doc.sections:
            # Header
            header = section.header
            for paragraph in header.paragraphs:
                if placeholder in paragraph.text:
                    placeholder_exists = True
                    print(f"Replacing '{placeholder}' with '{clean_value}' in header")
                    replace_text_with_custom_font(paragraph, placeholder, str(clean_value), preferred_font)
                    replacements_made += 1
            
            # Footer
            footer = section.footer
            for paragraph in footer.paragraphs:
                if placeholder in paragraph.text:
                    placeholder_exists = True
                    print(f"Replacing '{placeholder}' with '{clean_value}' in footer")
                    replace_text_with_custom_font(paragraph, placeholder, str(clean_value), preferred_font)
                    replacements_made += 1
        
        if not placeholder_exists and clean_value:
            print(f"WARNING: Placeholder '{placeholder}' not found in template!")
    
    print(f"Total replacements made: {replacements_made}")
    
    # Handle "Today" placeholder with Arabic date
    today = datetime.today()
    day = convert_to_eastern_arabic(today.day)
    month_name = arabic_months[today.strftime("%B")]
    year = convert_to_eastern_arabic(today.year)
    arabic_date = f"{day} {month_name} {year}"
    
    print(f"Replacing 'Today' with '{arabic_date}'")
    
    # Replace "Today" in all locations
    for paragraph in doc.paragraphs:
        if "Today" in paragraph.text:
            replace_text_with_custom_font(paragraph, "Today", arabic_date, preferred_font)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "Today" in paragraph.text:
                        replace_text_with_custom_font(paragraph, "Today", arabic_date, preferred_font)
    
    for section in doc.sections:
        # Header
        for paragraph in section.header.paragraphs:
            if "Today" in paragraph.text:
                replace_text_with_custom_font(paragraph, "Today", arabic_date, preferred_font)
        
        # Footer
        for paragraph in section.footer.paragraphs:
            if "Today" in paragraph.text:
                replace_text_with_custom_font(paragraph, "Today", arabic_date, preferred_font)
    
    # NUCLEAR OPTION: Completely rebuild document with consistent font
    nuclear_font_fix(doc, preferred_font)
    
    print("=== END TEMPLATE FILLING DEBUG ===")
    
    doc.save(output_path)
    return True, output_path
    
    print(f"Total replacements made: {replacements_made}")
    
    # Handle "Today" placeholder with Arabic date
    today = datetime.today()
    day = convert_to_eastern_arabic(today.day)
    month_name = arabic_months[today.strftime("%B")]
    year = convert_to_eastern_arabic(today.year)
    arabic_date = f"{day} {month_name} {year}"
    
    print(f"Replacing 'Today' with '{arabic_date}'")
    
    # Replace "Today" in all locations
    for paragraph in doc.paragraphs:
        if "Today" in paragraph.text:
            replace_text_with_custom_font(paragraph, "Today", arabic_date, preferred_font)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "Today" in paragraph.text:
                        replace_text_with_custom_font(paragraph, "Today", arabic_date, preferred_font)
    
    for section in doc.sections:
        # Header
        for paragraph in section.header.paragraphs:
            if "Today" in paragraph.text:
                replace_text_with_custom_font(paragraph, "Today", arabic_date, preferred_font)
        
        # Footer
        for paragraph in section.footer.paragraphs:
            if "Today" in paragraph.text:
                replace_text_with_custom_font(paragraph, "Today", arabic_date, preferred_font)
    
    # FINAL STEP: Apply consistent font to entire document
    apply_consistent_font_to_document(doc, preferred_font)
    
    print("=== END TEMPLATE FILLING DEBUG ===")
    
    doc.save(output_path)
    return True, output_path

def nuclear_font_fix(doc, font_name='Adobe Arabic'):
    """Nuclear option: completely rebuild document with consistent font"""
    print(f"NUCLEAR FONT FIX: Rebuilding entire document with '{font_name}'...")
    
    # Process every single paragraph
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Only process paragraphs with text
            force_paragraph_font(paragraph, font_name)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        force_paragraph_font(paragraph, font_name)
    
    # Process headers and footers
    for section in doc.sections:
        try:
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    force_paragraph_font(paragraph, font_name)
        except:
            pass
            
        try:
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip():
                    force_paragraph_font(paragraph, font_name)
        except:
            pass
    
    print(f"NUCLEAR FONT FIX COMPLETED with '{font_name}'")

def force_paragraph_font(paragraph, font_name='Adobe Arabic'):
    """Force a specific font on an entire paragraph while preserving original size"""
    # Get all text from the paragraph
    full_text = paragraph.text
    
    # Store paragraph-level formatting
    paragraph_format = paragraph.paragraph_format
    alignment = paragraph_format.alignment
    
    # Capture original font size from the first run with content
    original_size = None
    original_bold = None
    original_italic = None
    original_underline = None
    original_color = None
    
    for run in paragraph.runs:
        if run.text.strip():  # First run with actual text
            original_size = run.font.size
            original_bold = run.font.bold
            original_italic = run.font.italic
            original_underline = run.font.underline
            if run.font.color and run.font.color.rgb:
                original_color = run.font.color.rgb
            break
    
    # Clear all runs
    for run in paragraph.runs[:]:
        paragraph._element.remove(run._element)
    
    # Create one new run with the entire text
    if full_text.strip():  # Only if there's actual text
        new_run = paragraph.add_run(full_text)
        
        # Set font using multiple methods to ensure it sticks
        new_run.font.name = font_name
        new_run.font.cs_font = font_name  # Complex script (Arabic)
        new_run.font.east_asian_font = font_name  # East Asian
        
        # PRESERVE ORIGINAL FORMATTING
        if original_size:
            new_run.font.size = original_size
            print(f"Preserved font size: {original_size}")
        if original_bold is not None:
            new_run.font.bold = original_bold
        if original_italic is not None:
            new_run.font.italic = original_italic
        if original_underline is not None:
            new_run.font.underline = original_underline
        if original_color:
            new_run.font.color.rgb = original_color
        
        # Try setting the theme font as well
        try:
            new_run.font.theme_font = None  # Disable theme font
        except:
            pass
            
        # Set font using lower-level XML if available
        try:
            from docx.oxml.shared import qn
            rPr = new_run._element.get_or_add_rPr()
            # ASCII font
            ascii_font = rPr.find(qn('w:rFonts'))
            if ascii_font is None:
                ascii_font = rPr.add_child(qn('w:rFonts'))
            ascii_font.set(qn('w:ascii'), font_name)
            ascii_font.set(qn('w:hAnsi'), font_name)
            ascii_font.set(qn('w:cs'), font_name)  # Complex script
            ascii_font.set(qn('w:eastAsia'), font_name)  # East Asian
        except Exception as e:
            print(f"Could not set low-level font: {e}")
        
        # Restore paragraph alignment
        paragraph_format.alignment = alignment
        
        size_info = f" (size: {original_size})" if original_size else ""
        print(f"Forced font '{font_name}'{size_info} on paragraph: '{full_text[:50]}...'")

def replace_text_with_custom_font(paragraph, old_text, new_text, font_name='Adobe Arabic'):
    """Replace text with a specific font and ensure consistent formatting while preserving size"""
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return
    
    # Store the original paragraph formatting
    paragraph_format = paragraph.paragraph_format
    alignment = paragraph_format.alignment
    
    # Get the first run's formatting as baseline - IMPROVED VERSION
    baseline_run = None
    for run in paragraph.runs:
        if run.text.strip():
            baseline_run = run
            break
    
    if not baseline_run and paragraph.runs:
        baseline_run = paragraph.runs[0]
    
    # Store formatting properties with better defaults
    if baseline_run:
        original_size = baseline_run.font.size
        original_bold = baseline_run.font.bold
        original_italic = baseline_run.font.italic
        original_underline = baseline_run.font.underline
        original_color = baseline_run.font.color.rgb if baseline_run.font.color and baseline_run.font.color.rgb else None
    else:
        # Better default formatting - try to get from paragraph style
        from docx.shared import Pt
        original_size = Pt(12)  # Default size
        original_bold = False
        original_italic = False
        original_underline = False
        original_color = None
        
        # Try to get size from paragraph style
        try:
            if paragraph.style and paragraph.style.font and paragraph.style.font.size:
                original_size = paragraph.style.font.size
                print(f"Using paragraph style font size: {original_size}")
        except:
            pass
    
    # Replace text
    full_text = full_text.replace(old_text, new_text)
    
    # Clear all runs
    for run in paragraph.runs[:]:
        paragraph._element.remove(run._element)
    
    # Create a single new run with consistent formatting
    new_run = paragraph.add_run(full_text)
    
    # Apply font using multiple methods
    new_run.font.name = font_name
    new_run.font.cs_font = font_name  # Complex script font for Arabic
    new_run.font.east_asian_font = font_name  # East Asian font
    
    # PRESERVE ORIGINAL FORMATTING - this is the key fix
    if original_size:
        new_run.font.size = original_size
        print(f"Preserved original font size: {original_size} for text: '{new_text[:30]}...'")
    if original_bold is not None:
        new_run.font.bold = original_bold
    if original_italic is not None:
        new_run.font.italic = original_italic
    if original_underline is not None:
        new_run.font.underline = original_underline
    if original_color:
        new_run.font.color.rgb = original_color
    
    # Try setting theme font to None
    try:
        new_run.font.theme_font = None
    except:
        pass
    
    # Set font using lower-level XML
    try:
        from docx.oxml.shared import qn
        rPr = new_run._element.get_or_add_rPr()
        ascii_font = rPr.find(qn('w:rFonts'))
        if ascii_font is None:
            ascii_font = rPr.add_child(qn('w:rFonts'))
        ascii_font.set(qn('w:ascii'), font_name)
        ascii_font.set(qn('w:hAnsi'), font_name)
        ascii_font.set(qn('w:cs'), font_name)
        ascii_font.set(qn('w:eastAsia'), font_name)
    except Exception as e:
        print(f"Could not set low-level font: {e}")
    
    # Restore alignment
    paragraph_format.alignment = alignment
    
    size_info = f" (size: {original_size})" if original_size else ""
    print(f"Applied font '{font_name}'{size_info} to: '{new_text[:50]}...'")

# Keep Adobe Arabic as the preferred font
PREFERRED_ARABIC_FONT = 'Adobe Arabic'

def load_data():
    try:
        df = pd.read_excel(app.config['EXCEL_FILE'], header=1)
        df = df.loc[:, ~df.columns.str.contains('^Unnamed')]
        
        # Clean column names by removing leading/trailing spaces
        df.columns = df.columns.str.strip()
        
        # Handle NaN/None values first
        df = df.fillna('')
        
        # Clean date columns before converting to strings
        df = clean_dataframe_dates(df)
        
        # Convert all values to strings
        df = df.astype(str)
        return df
    except Exception as e:
        app.logger.error(f"Error loading Excel file: {e}")
        return pd.DataFrame()

def get_gender_value(row):
    """Extracts and normalizes gender value from row"""
    gender = str(row.get('Gender', row.get('gender', ''))).strip().upper()
    return gender in ('M', 'MALE', '1', 'ذكر')  # Returns True for male, False for female

@app.route('/', methods=['GET', 'POST'])
def index():
    df = load_data()
    search_term = request.form.get('search', '').strip().lower()
    
    if not df.empty:
        if search_term:
            filtered_rows = df.apply(
                lambda row: row.str.contains(search_term, case=False).any(),
                axis=1
            )
            rows = df[filtered_rows].to_dict('records')
        else:
            rows = df.to_dict('records')
    else:
        rows = []
        flash("No data available or error loading data", "warning")
    
    return render_template('index.html', rows=rows, search_term=search_term)

@app.route('/generate', methods=['POST'])
def generate_document():
    try:
        row_index = int(request.form.get('row_index'))
        doc_type = request.form.get('doc_type')
        
        df = load_data()
        if df.empty:
            flash("Error loading Excel file", "error")
            return redirect(url_for('index'))
        
        selected_row = df.iloc[row_index]
        is_male = get_gender_value(selected_row)
        
        # Determine template and filename
        if doc_type == 'baptism':
            template = 'baptisim_template_m.docx' if is_male else 'baptisim_template_f.docx'
            filename = f"معمودية{row_index + 1}.docx"
        else:  # release
            template = 'release_situation_m.docx' if is_male else 'release_situation_f.docx'
            filename = f"اطلاق حال{row_index + 1}.docx"
        
        template_path = os.path.join(app.config['TEMPLATE_DIR'], template)
        output_path = os.path.join(app.config['OUTPUT_DIR'], filename)
        
        # Verify template exists
        if not os.path.exists(template_path):
            flash(f"Template file not found: {template}", "error")
            return redirect(url_for('index'))
        
        success, message = fill_template(template_path, selected_row, output_path)
        
        if success:
#            flash(f"تم إنشاء الملف: {filename}", "success")
            return send_from_directory(app.config['OUTPUT_DIR'], filename, as_attachment=True)
        else:
            flash(message, "error")
            return redirect(url_for('index'))
    
    except Exception as e:
        flash(f"An error occurred: {str(e)}", "error")
        return redirect(url_for('index'))

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(app.config['OUTPUT_DIR'], filename, as_attachment=True)

@app.route('/debug')
def debug_data():
    try:
        df = load_data()
        if df.empty:
            return {"status": "error", "message": "DataFrame is empty"}
        
        first_row = df.iloc[0].to_dict()
        return {
            "status": "success",
            "columns": list(df.columns),
            "first_row": first_row,
            "gender_info": {
                "raw_value": first_row.get('Gender', 'N/A'),
                "normalized": get_gender_value(df.iloc[0])
            }
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}

if __name__ == '__main__':
    # Create necessary directories if they don't exist
    os.makedirs(app.config['OUTPUT_DIR'], exist_ok=True)
    os.makedirs(app.config['TEMPLATE_DIR'], exist_ok=True)
    os.makedirs('data', exist_ok=True)
    
    # Get port from environment variable
    port = int(os.environ.get('PORT', 5000))
    
    # Run the app - MUST use host='0.0.0.0' for Render
    app.run(host='0.0.0.0', port=port, debug=False)